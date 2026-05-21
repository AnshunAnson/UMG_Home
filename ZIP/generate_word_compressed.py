#!/usr/bin/env python3
import json
import os
import subprocess
import zipfile
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import tempfile

FFMPEG = "C:\\Program Files\\UI2V\\resources\\app.asar.unpacked\\node_modules\\ffmpeg-static\\ffmpeg.exe"
IMG_WIDTH = Inches(4.5)

def load_content():
    with open('content.json', 'r', encoding='utf-8') as f:
        return json.load(f)

def get_media_path(src):
    base_paths = ['images/', 'gifs/']
    for base in base_paths:
        full_path = base + src.replace('images/', '').replace('gifs/', '')
        if os.path.exists(full_path):
            return full_path
    if os.path.exists(src):
        return src
    return None

def is_video(src):
    ext = src.lower().split('.')[-1]
    return ext in ['mp4', 'webm', 'avi', 'wmv']

def video_to_compressed_gif(video_path, output_gif, max_width=640):
    try:
        subprocess.run([
            FFMPEG, '-y', '-i', video_path,
            '-vf', f'fps=10,scale={max_width}:-1:flags=lanczos,split[s0][s1];[s0]palettegen=max_colors=64[p];[s1][p]paletteuse=dither=floyd_steinberg',
            '-loop', '0', output_gif
        ], check=True, capture_output=True, text=True)
        print(f"  Created compressed GIF: {output_gif} ({os.path.getsize(output_gif) // 1024} KB)")
        return output_gif
    except Exception as e:
        print(f"  Error converting {video_path}: {e}")
        if os.path.exists(output_gif):
            os.remove(output_gif)
        return None

def compress_existing_gif(gif_path, output_gif, max_width=640):
    try:
        subprocess.run([
            FFMPEG, '-y', '-i', gif_path,
            '-vf', f'fps=10,scale={max_width}:-1:flags=lanczos,split[s0][s1];[s0]palettegen=max_colors=64[p];[s1][p]paletteuse=dither=floyd_steinberg',
            '-loop', '0', output_gif
        ], check=True, capture_output=True, text=True)
        print(f"  Compressed GIF: {output_gif} ({os.path.getsize(output_gif) // 1024} KB)")
        return output_gif
    except Exception as e:
        print(f"  Error compressing {gif_path}: {e}")
        return gif_path

def embed_gif_in_docx(doc_path, gif_path, description=""):
    gif_filename = os.path.basename(gif_path)
    media_target = f'word/media/{gif_filename}'

    with zipfile.ZipFile(doc_path, 'r') as zf:
        if media_target in zf.namelist():
            print(f"  GIF {gif_filename} already embedded")
            return True

    with open(gif_path, 'rb') as f:
        gif_data = f.read()

    tmp_path = tempfile.mktemp(suffix='.docx')
    shutil.copy(doc_path, tmp_path)

    with zipfile.ZipFile(tmp_path, 'r') as zf_in:
        with zipfile.ZipFile(doc_path, 'w', zipfile.ZIP_DEFLATED) as zf_out:
            for item in zf_in.namelist():
                if item == 'word/_rels/document.xml.rels':
                    rels_content = zf_in.read(item).decode('utf-8')
                    import re
                    ids = re.findall(r'Id="rId(\d+)"', rels_content)
                    existing_id = max([int(x) for x in ids]) + 1 if ids else 1

                    if media_target not in rels_content:
                        new_rel = f'<Relationship Id="rId{existing_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/{gif_filename}"/>'
                        rels_content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')
                    zf_out.writestr(item, rels_content.encode('utf-8'))
                elif item == 'word/document.xml':
                    doc_content = zf_in.read(item).decode('utf-8')
                    zf_out.writestr(item, doc_content.encode('utf-8'))
                else:
                    zf_out.writestr(item, zf_in.read(item))
            zf_out.writestr(media_target, gif_data)

    os.unlink(tmp_path)
    print(f"  Embedded GIF: {gif_filename}")
    return True

def create_word_with_compressed_gifs(content, output_path='portfolio_compressed.docx'):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'SimHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    style.font.size = Pt(22)

    hero = content['heroContent']
    about = content['aboutContent']
    projects = content['projectsContent']['projects']
    skills = content['skillsContent']
    contact = content['contactContent']

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("https://anshunanson.github.io/Personal_Technical_Homepage/")
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hero['badge'])
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(128, 128, 128)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hero['name'])
    run.font.size = Pt(60)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(hero['subtitle'])
    run.font.size = Pt(24)

    doc.add_paragraph()

    info = doc.add_paragraph()
    run = info.add_run(f"职位: {about['jobTitle']}")
    run.font.size = Pt(20)
    info.add_run("\n")
    run = info.add_run(f"工作年限: {about['experience']} 年")
    run.font.size = Pt(20)
    info.add_run("\n")
    run = info.add_run(f"年龄: {about['age']}")
    run.font.size = Pt(20)

    doc.add_paragraph()
    for bio in about['bio']:
        p = doc.add_paragraph(bio)
        p.runs[0].font.size = Pt(20)

    doc.add_paragraph()

    heading = doc.add_heading('项目作品', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 102)

    gif_map = {}
    for project in projects:
        p = doc.add_heading(project['title'], level=2)

        info_parts = []
        if project.get('period') and project['period'] != 'Unspecified':
            info_parts.append(project['period'])
        if project.get('category'):
            info_parts.append(project['category'])
        if info_parts:
            p = doc.add_paragraph(' | '.join(info_parts))
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        if project.get('description'):
            p = doc.add_paragraph(project['description'])
            p.runs[0].font.size = Pt(22)

        if project.get('tech'):
            p = doc.add_paragraph(f"技术栈: {' | '.join(project['tech'])}")
            p.runs[0].font.size = Pt(18)
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        if project.get('details'):
            for detail in project['details']:
                p = doc.add_paragraph(f"- {detail}")
                p.runs[0].font.size = Pt(20)

        if project.get('subProjects'):
            for sub in project['subProjects']:
                p = doc.add_paragraph()
                run = p.add_run(f">> {sub['title']}")
                run.font.bold = True
                run.font.size = Pt(24)

                if sub.get('description'):
                    p = doc.add_paragraph(sub['description'])
                    p.runs[0].font.size = Pt(20)

                if sub.get('images'):
                    for i, img in enumerate(sub['images']):
                        if i >= 2:
                            break
                        media_path = get_media_path(img['src'])
                        if media_path and os.path.exists(media_path):
                            if is_video(img['src']):
                                gif_filename = os.path.basename(media_path).rsplit('.', 1)[0] + '_compressed.gif'
                                gif_path = os.path.join('gifs', gif_filename)
                                if not os.path.exists(gif_path):
                                    print(f"Converting video to compressed GIF: {media_path}")
                                    gif_path = video_to_compressed_gif(media_path, gif_path)
                                if gif_path and os.path.exists(gif_path):
                                    gif_map[gif_path] = img.get('alt', gif_filename)
                                    try:
                                        doc.add_picture(gif_path, width=IMG_WIDTH)
                                        if img.get('alt'):
                                            p = doc.add_paragraph(img['alt'])
                                            p.runs[0].font.size = Pt(18)
                                    except Exception as e:
                                        print(f"Warning: Cannot embed GIF {gif_path}: {e}")
                            else:
                                try:
                                    doc.add_picture(media_path, width=IMG_WIDTH)
                                    alt = img.get('alt')
                                    if alt:
                                        p = doc.add_paragraph(alt)
                                        p.runs[0].font.size = Pt(18)
                                except Exception as e:
                                    print(f"Warning: Cannot embed {media_path}: {e}")

        if project.get('images'):
            for i, img in enumerate(project['images']):
                if i >= 3:
                    break
                media_path = get_media_path(img['src'])
                if media_path and os.path.exists(media_path):
                    if is_video(img['src']):
                        gif_filename = os.path.basename(media_path).rsplit('.', 1)[0] + '_compressed.gif'
                        gif_path = os.path.join('gifs', gif_filename)
                        if not os.path.exists(gif_path):
                            print(f"Converting video to compressed GIF: {media_path}")
                            gif_path = video_to_compressed_gif(media_path, gif_path)
                        if gif_path and os.path.exists(gif_path):
                            gif_map[gif_path] = img.get('alt', gif_filename)
                            try:
                                doc.add_picture(gif_path, width=IMG_WIDTH)
                                if img.get('alt'):
                                    p = doc.add_paragraph(img['alt'])
                                    p.runs[0].font.size = Pt(18)
                            except Exception as e:
                                print(f"Warning: Cannot embed GIF {gif_path}: {e}")
                    else:
                        try:
                            doc.add_picture(media_path, width=IMG_WIDTH)
                            alt = img.get('alt')
                            if alt:
                                p = doc.add_paragraph(alt)
                                p.runs[0].font.size = Pt(18)
                        except Exception as e:
                            print(f"Warning: Cannot embed {media_path}: {e}")

        if project.get('links'):
            for link in project['links']:
                p = doc.add_paragraph(f"链接: {link.get('label', '查看')} - {link.get('href', '')}")
                p.runs[0].font.color.rgb = RGBColor(0, 102, 153)
                p.runs[0].font.size = Pt(20)

        doc.add_paragraph()

    heading = doc.add_heading('技术能力', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 102)

    for category in skills['categories']:
        p = doc.add_heading(category['title'], level=2)
        p = doc.add_paragraph(' | '.join([s['name'] for s in category['skills']]))
        p.runs[0].font.size = Pt(20)

    p = doc.add_heading('技术栈概览', level=2)
    p = doc.add_paragraph(' | '.join(skills['techStack']))
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    heading = doc.add_heading('联系方式', level=1)
    heading.runs[0].font.color.rgb = RGBColor(0, 102, 102)

    p = doc.add_paragraph(contact['description'])
    p.runs[0].font.size = Pt(20)
    p = doc.add_paragraph(f"邮箱: {contact['email']}")
    p.runs[0].font.size = Pt(20)
    p = doc.add_paragraph(f"电话: {contact['phone']}")
    p.runs[0].font.size = Pt(20)
    p = doc.add_paragraph("Website: https://anshunanson.github.io/Personal_Technical_Homepage/")
    p.runs[0].font.color.rgb = RGBColor(0, 102, 153)
    p.runs[0].font.size = Pt(20)

    doc.save(output_path)
    print(f"\nDocument saved: {output_path}")
    print(f"Embedding {len(gif_map)} compressed GIFs...")

    for gif_path, description in gif_map.items():
        if os.path.exists(gif_path):
            embed_gif_in_docx(output_path, gif_path, description)

    print(f"\nWord document with compressed content generated: {output_path}")

if __name__ == '__main__':
    content = load_content()
    create_word_with_compressed_gifs(content)
